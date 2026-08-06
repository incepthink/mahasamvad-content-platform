from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("government-influencer-campaign-platform-concept.docx")

FONT = "Calibri"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2933"
MUTED = "59636E"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "C7CED6"

TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM = 80
CELL_START_END = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, *, name: str = FONT, size: float | None = None,
                 color: str | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, *, top: int = CELL_TOP_BOTTOM,
                     start: int = CELL_START_END,
                     bottom: int = CELL_TOP_BOTTOM,
                     end: int = CELL_START_END) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {TABLE_WIDTH_DXA}: {widths_dxa}")

    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color: str = BORDER, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_padding(paragraph, *, before: int = 90, after: int = 90,
                          left: int = 140, right: int = 140) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:right"), str(right))
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))


def add_page_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_numbering_definitions(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element

    existing_abs = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs, default=0) + 1
    next_num = max(existing_num, default=0) + 1

    def abstract(abstract_id: int, num_fmt: str, level_text: str, font: str | None = None):
        abs_num = OxmlElement("w:abstractNum")
        abs_num.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abs_num.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl.append(fmt)
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), level_text)
        lvl.append(text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)

        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)

        r_pr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), font or FONT)
        rfonts.set(qn("w:hAnsi"), font or FONT)
        rfonts.set(qn("w:eastAsia"), font or FONT)
        r_pr.append(rfonts)
        lvl.append(r_pr)
        abs_num.append(lvl)
        numbering.append(abs_num)

    def instance(num_id: int, abstract_id: int):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    bullet_abs = next_abs
    decimal_abs = next_abs + 1
    bullet_num = next_num
    decimal_num = next_num + 1
    abstract(bullet_abs, "bullet", "\u2022", FONT)
    abstract(decimal_abs, "decimal", "%1.", FONT)
    instance(bullet_num, bullet_abs)
    instance(decimal_num, decimal_abs)
    return bullet_num, decimal_num


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_bullet(doc: Document, text: str, bullet_num: int, *, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    apply_numbering(p, bullet_num)
    if bold_prefix and text.startswith(bold_prefix):
        lead = p.add_run(bold_prefix)
        set_run_font(lead, bold=True, color=INK)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, color=INK)
    return p


def add_step(doc: Document, title: str, body: str, decimal_num: int, *, output: str | None = None):
    p = doc.add_paragraph()
    apply_numbering(p, decimal_num)
    p.paragraph_format.keep_with_next = True
    title_run = p.add_run(title)
    set_run_font(title_run, bold=True, color=NAVY)

    detail = doc.add_paragraph(body)
    detail.paragraph_format.left_indent = Inches(0.5)
    detail.paragraph_format.first_line_indent = Inches(0)
    detail.paragraph_format.space_after = Pt(5)
    detail.paragraph_format.line_spacing = 1.10
    detail.paragraph_format.keep_together = True
    for run in detail.runs:
        set_run_font(run, color=INK)

    if output:
        out = doc.add_paragraph()
        out.paragraph_format.left_indent = Inches(0.5)
        out.paragraph_format.space_after = Pt(8)
        out.paragraph_format.keep_together = True
        label = out.add_run("Result: ")
        set_run_font(label, bold=True, color=DARK_BLUE)
        value = out.add_run(output)
        set_run_font(value, color=INK)


def add_label_paragraph(
    doc: Document,
    label: str,
    value: str,
    *,
    keep: bool = True,
    keep_with_next: bool = False,
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_together = keep
    p.paragraph_format.keep_with_next = keep_with_next
    lead = p.add_run(f"{label}: ")
    set_run_font(lead, bold=True, color=DARK_BLUE)
    rest = p.add_run(value)
    set_run_font(rest, color=INK)
    return p


def add_note(doc: Document, label: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.10
    set_paragraph_shading(p, CALLOUT)
    set_paragraph_padding(p)
    lead = p.add_run(f"{label}: ")
    set_run_font(lead, bold=True, color=DARK_BLUE)
    body = p.add_run(text)
    set_run_font(body, color=INK)
    return p


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def configure_header_footer(section) -> None:
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("Government Influencer Campaign Management Platform")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    prefix = fp.add_run("Concept Note  |  ")
    set_run_font(prefix, size=9, color=MUTED)
    add_page_field(fp, "PAGE")
    divider = fp.add_run(" of ")
    set_run_font(divider, size=9, color=MUTED)
    add_page_field(fp, "NUMPAGES")


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(6)
    kicker = p.add_run("PRODUCT CONCEPT NOTE")
    set_run_font(kicker, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("Government Influencer\nCampaign Management Platform")
    set_run_font(run, size=27, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(28)
    subrun = subtitle.add_run("Operational flow, platform resources and example task formats")
    set_run_font(subrun, size=13, color=MUTED)

    metadata = [
        ("Document", "Product concept and operational flow"),
        ("Scope", "Influencer registration through submission evaluation, tracking and payment"),
        ("Status", "Working concept note"),
        ("Date", "6 August 2026"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, color=NAVY, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.5, color=INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(38)
    add_note(
        doc,
        "Scope of this note",
        "This document describes the user roles, step-by-step operating flow, proof-of-work approach, high-level hosting and service resources, and representative task formats. The payment method, hosting vendor and external analytics providers remain open for government selection.",
    )
    doc.add_page_break()


def add_platform_overview(doc: Document, bullet_num: int) -> None:
    doc.add_heading("1. Platform Overview", level=1)
    p = doc.add_paragraph(
        "The platform will allow government departments to create campaigns, define campaign-specific eligibility, select influencers, assign tasks, communicate with participants, evaluate proof-of-work submissions, approve payment records and monitor campaign progress."
    )
    p.paragraph_format.keep_together = True

    add_note(
        doc,
        "Operating principle",
        "Influencers maintain one reusable base profile. Additional evidence is collected only when a campaign requires it. A regional campaign may request residence, audience-location or previous-work evidence; a PAN-India campaign may skip regional verification.",
    )

    doc.add_heading("1.1 Two connected interfaces", level=2)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [4680, 4680])
    set_table_borders(table)
    headers = ["Influencer portal", "Government dashboard"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, bold=True, color=NAVY)
    repeat_table_header(table.rows[0])
    row = table.add_row()
    for cell in row.cells:
        keep_row_together(row)
    portal = (
        "Registration and profile management; campaign discovery and applications; eligibility evidence; assigned tasks; draft approval where required; campaign chat; proof-of-work submission; correction requests; and payment-status visibility."
    )
    dashboard = (
        "Campaign and task creation; eligibility rules; influencer review and selection; regional verification; communication; submission evaluation; payment approval and status; campaign monitoring; reporting and export."
    )
    for idx, value in enumerate((portal, dashboard)):
        p = row.cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=10.2, color=INK)

    doc.add_heading("1.2 Platform and hosting resources", level=2)
    p = doc.add_paragraph(
        "The following capabilities must be provisioned before launch. Exact capacity and vendor selection should be based on the number of active influencers, concurrent campaigns, expected file sizes, peak campaign traffic and the government's retention policy."
    )
    p.paragraph_format.keep_with_next = True

    resources = [
        ("Application environments", "Runs the influencer portal, government dashboard and a safe testing environment.", "Separate production and staging environments; production scales for peak campaign activity."),
        ("Database service", "Stores profiles, campaigns, tasks, applications, decisions and payment-status records.", "Managed, encrypted and restricted to authorised users and services."),
        ("File and media storage", "Stores residence evidence, drafts, screenshots, videos, analytics proof and chat attachments.", "Capacity based on submissions multiplied by average file size and retention period."),
        ("Identity and access", "Supports verified sign-in and role-based access for influencers and government officers.", "Mobile OTP or email verification; stronger access controls for government users."),
        ("Chat and notifications", "Supports campaign communication and deadline, correction, approval and payment alerts.", "In-platform chat plus selected email, SMS or approved messaging channels."),
        ("Security, monitoring and recovery", "Protects the service, records important actions, detects outages and restores lost data.", "Encryption, certificates, web protection, file scanning, alerts, audit logs, automated backups and tested recovery."),
        ("Integration and reporting capacity", "Connects to payment, reporting or social analytics tools selected later.", "Secure import, export or integration capability without locking the product to one provider."),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2150, 3220, 3990])
    set_table_borders(table)
    for idx, text in enumerate(("Resource", "What it supports", "Planning consideration")):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=9.5, bold=True, color=NAVY)
    repeat_table_header(table.rows[0])
    for resource, purpose, planning in resources:
        row = table.add_row()
        keep_row_together(row)
        values = (resource, purpose, planning)
        for idx, value in enumerate(values):
            p = row.cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=9.2, color=INK, bold=(idx == 0))

    doc.add_page_break()


def add_users(doc: Document) -> None:
    doc.add_heading("2. Users and Responsibilities", level=1)
    p = doc.add_paragraph(
        "The platform separates responsibilities so that campaign decisions, regional verification, submission review and payment approval can be controlled and audited. In a small pilot, one authorised officer may perform more than one role."
    )
    p.paragraph_format.keep_with_next = True

    roles = [
        ("Influencer", "Maintains a profile, applies to eligible campaigns, completes assigned tasks, communicates with officers and submits proof of work."),
        ("Campaign manager", "Creates campaigns and tasks, defines eligibility, selects influencers, manages deadlines and monitors delivery."),
        ("Regional verification officer", "Checks residence, local audience, language ability or previous regional work when the campaign requires regional eligibility."),
        ("Content or communication officer", "Provides approved messages, checks drafts when pre-publication approval is required and guards factual accuracy."),
        ("Submission evaluator", "Reviews proof of work, applies the task criteria and accepts, rejects or requests correction with a recorded reason."),
        ("Finance officer", "Reviews accepted work, confirms the payable amount and updates or imports the payment status from the selected payment process."),
        ("Platform administrator", "Manages officer access, role permissions, platform settings, escalation support and audit access."),
        ("Reporting user", "Views campaign progress, regional participation, task completion, performance and payment summaries without changing operational records."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2500, 6860])
    set_table_borders(table)
    for idx, text in enumerate(("User or role", "Primary responsibility")):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, bold=True, color=NAVY)
    repeat_table_header(table.rows[0])
    for role, responsibility in roles:
        row = table.add_row()
        keep_row_together(row)
        for idx, value in enumerate((role, responsibility)):
            p = row.cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=10.1, color=INK, bold=(idx == 0))

    doc.add_heading("2.1 Chat and formal decisions", level=2)
    p = doc.add_paragraph(
        "Government officers and influencers can chat within the relevant campaign or task. Chat may be used for questions, clarifications, attachments and reminders. Formal actions such as selection, task assignment, deadline extension, rejection, acceptance and payment approval must still be recorded through the relevant dashboard action so they remain visible in the audit history."
    )
    p.paragraph_format.keep_together = True
    doc.add_page_break()


def add_flow(doc: Document, decimal_num: int) -> None:
    doc.add_heading("3. Step-by-Step Product Flow", level=1)
    intro = doc.add_paragraph(
        "The flow begins with a reusable influencer account and ends with an evaluated submission, payment tracking and campaign reporting. Campaign rules determine which verification and proof fields appear at each stage."
    )
    intro.paragraph_format.keep_together = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    set_paragraph_shading(p, LIGHT_BLUE)
    set_paragraph_padding(p, before=110, after=110)
    run = p.add_run("REGISTER  >  APPLY  >  VERIFY  >  ASSIGN  >  DELIVER  >  SUBMIT  >  EVALUATE  >  APPROVE  >  TRACK")
    set_run_font(run, size=10, color=NAVY, bold=True)

    doc.add_heading("Stage A - Registration and profile", level=2)
    add_step(
        doc,
        "Influencer registration",
        "The influencer creates an account using a verified mobile number or email address and accepts the platform terms, privacy notice and campaign participation rules.",
        decimal_num,
        output="A verified account that can be used across multiple campaigns.",
    )
    add_step(
        doc,
        "Reusable influencer profile",
        "The influencer records basic identity, languages, content categories, social handles, audience size, primary audience locations, regions where physical work is possible and previous campaign experience. Sensitive evidence is not required at this stage unless government policy makes it mandatory.",
        decimal_num,
        output="A reusable profile that reduces repeated form filling and supports campaign screening.",
    )

    doc.add_heading("Stage B - Campaign eligibility and selection", level=2)
    add_step(
        doc,
        "Government campaign creation",
        "The campaign manager creates a campaign with its objective, geographic scope, target audience, timeline, budget context, number of influencers, eligibility rules, tasks, proof requirements and evaluation conditions.",
        decimal_num,
        output="A campaign ready to publish to eligible influencers or use for direct invitations.",
    )
    add_step(
        doc,
        "Campaign discovery or invitation",
        "Influencers browse available campaigns or receive an invitation. Before applying, they can see the region, eligibility, deliverables, deadlines, proof requirements and payment basis.",
        decimal_num,
        output="An informed application or invitation response.",
    )
    add_step(
        doc,
        "Campaign-specific eligibility and verification",
        "The platform asks only for evidence required by that campaign. A regional campaign may request residence evidence, primary audience location, local language ability or previous regional work. A PAN-India campaign may skip regional evidence and use broader criteria such as category, reach, platform or language.",
        decimal_num,
        output="An application marked eligible, ineligible or requiring officer review.",
    )
    add_step(
        doc,
        "Government review and selection",
        "Authorised officers review the application and supporting evidence. They may approve, reject with a reason, request more information, shortlist the influencer or place the applicant on a waiting list. Eligibility does not guarantee selection when campaign slots are limited.",
        decimal_num,
        output="A recorded campaign participation decision.",
    )

    doc.add_heading("Stage C - Task assignment and delivery", level=2)
    add_step(
        doc,
        "Task assignment",
        "An approved influencer receives one or more task briefs. Each brief specifies the platform, content format, mandatory messages, language, audience, deadlines, prohibited claims, proof fields, evaluation criteria and payment rule.",
        decimal_num,
        output="A clear, accepted assignment linked to the campaign and influencer.",
    )
    add_step(
        doc,
        "Optional pre-publication approval",
        "When a campaign is sensitive, the influencer submits a draft before publishing. The communication officer approves it or requests changes. For lower-risk tasks, the campaign may permit direct publication followed by proof submission.",
        decimal_num,
        output="An approved draft or permission to proceed under the direct-publication route.",
    )
    add_step(
        doc,
        "Work execution and campaign chat",
        "The influencer completes the assigned activity and uses task-linked chat to ask questions, share files or request clarification. Officers can issue reminders and respond without moving the conversation to an unrecorded channel.",
        decimal_num,
        output="Completed campaign activity with a traceable communication history.",
    )

    doc.add_heading("Stage D - Proof of work and evaluation", level=2)
    add_step(
        doc,
        "Proof-of-work submission",
        "The submission form changes according to the task. It may request a public URL, screenshot, original media, analytics evidence, event photograph, location evidence, attendance confirmation or completion note. Every submission is timestamped and linked to the specific assignment.",
        decimal_num,
        output="A complete evidence package ready for evaluation.",
    )
    add_step(
        doc,
        "Submission evaluation",
        "The evaluator checks whether the work followed the brief, used accurate and appropriate messaging, met the deadline and included valid proof. Performance information such as reach and engagement can be recorded separately from basic task compliance.",
        decimal_num,
        output="Accepted, correction requested or rejected, with a recorded reason.",
    )
    add_step(
        doc,
        "Correction and resubmission",
        "If a correctable issue is found, the evaluator sends a specific correction request with a deadline. The influencer uploads a revised submission, while the earlier version and decision remain available in the history.",
        decimal_num,
        output="A final submission version and a traceable resolution.",
    )

    doc.add_heading("Stage E - Acceptance, payment and reporting", level=2)
    add_step(
        doc,
        "Final acceptance and payment eligibility",
        "When the task is accepted, the system creates an approved-for-payment record containing the influencer, campaign, task, accepted proof, payable amount, approving officer and approval date.",
        decimal_num,
        output="An auditable payment instruction ready for the selected payment process.",
    )
    add_step(
        doc,
        "Payment processing and status",
        "The government may pay through the dashboard or a separate approved tool. The platform records the approved amount, payment status, transaction reference, date and any hold or failure reason. The payment method remains configurable.",
        decimal_num,
        output="A payment record shown as approved, processing, paid, failed or on hold.",
    )
    add_step(
        doc,
        "Campaign monitoring and closure",
        "The dashboard summarises applications, selected influencers, assigned and completed tasks, pending evaluations, regional participation, campaign reach, accepted amounts and payment status. At closure, the final records remain available according to the government's retention policy.",
        decimal_num,
        output="A campaign summary and reusable influencer performance history.",
    )

    doc.add_heading("3.1 Main status journeys", level=2)
    statuses = [
        ("Application", "Available > Applied > Verification > Approved / Waitlisted / Rejected"),
        ("Task", "Assigned > Draft review if required > In progress > Proof submitted > Correction requested / Accepted / Rejected"),
        ("Payment", "Approved for payment > Processing > Paid / Failed / On hold"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2050, 7310])
    set_table_borders(table)
    for idx, text in enumerate(("Journey", "Representative statuses")):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, bold=True, color=NAVY)
    repeat_table_header(table.rows[0])
    for journey, status in statuses:
        row = table.add_row()
        keep_row_together(row)
        for idx, value in enumerate((journey, status)):
            p = row.cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=10.2, color=INK, bold=(idx == 0))
    doc.add_page_break()


def add_task_examples(doc: Document, bullet_num: int) -> None:
    doc.add_heading("4. Example Task Formats", level=1)
    p = doc.add_paragraph(
        "A campaign may contain several tasks, and each task defines its own proof-of-work form. The examples below show the level of clarity required in a task brief; the final wording, duration and payment rule remain campaign decisions."
    )
    p.paragraph_format.keep_together = True

    doc.add_heading("4.1 Standard task brief", level=2)
    for item in (
        "Task title and campaign objective",
        "Required platform, format, language and target audience",
        "Mandatory messages, approved references and prohibited claims",
        "Draft-approval requirement, publication window and final deadline",
        "Exact proof fields and the time at which analytics should be captured",
        "Evaluation criteria, correction limit and payment basis",
        "Required government-sponsored-content or paid-partnership disclosure",
    ):
        add_bullet(doc, item, bullet_num)

    add_note(
        doc,
        "Proof configuration",
        "The campaign manager chooses the proof fields while creating the task. Influencers see those fields before accepting the assignment, and a submission cannot be marked complete until all mandatory proof is provided.",
    )

    examples = [
        {
            "title": "4.2 Instagram or short-video reel",
            "deliverable": "Create and publish one short video using the approved campaign message, required language, campaign tags and disclosure. The task may specify a duration range and whether a draft must be approved first.",
            "proof": "Public post URL, publication screenshot, original or final video file, and analytics screenshot captured after the campaign-defined reporting period.",
            "evaluation": "Published on time; account ownership confirmed; mandatory message and disclosure present; prohibited claims absent; link accessible; proof legible.",
            "payment": "Normally a fixed amount after acceptance, unless the campaign explicitly adds a performance component.",
        },
        {
            "title": "4.3 Regional event attendance and coverage",
            "deliverable": "Attend a specified government event and publish the required coverage, such as one live update and one recap post. The campaign may restrict eligibility to influencers who live in, serve or have a verified audience in the relevant region.",
            "proof": "Officer or attendance confirmation, event photograph, published-content URL, timestamp, and location evidence only when the campaign has declared it necessary.",
            "evaluation": "Attendance verified; correct event represented; required coverage published; regional eligibility satisfied; government message accurately conveyed.",
            "payment": "May combine an attendance component and an accepted-content component, both defined before assignment.",
        },
        {
            "title": "4.4 YouTube or long-form explainer",
            "deliverable": "Produce a longer video explaining a scheme, public service or campaign. Because the format carries more factual detail, draft script or video approval may be mandatory before publication.",
            "proof": "Approved draft reference, public video URL, final uploaded file if required, thumbnail screenshot and analytics after the reporting period.",
            "evaluation": "Factual accuracy; use of approved terminology; required disclosures; clarity; acceptable audio and video quality; publication and proof deadlines met.",
            "payment": "A fixed deliverable amount or defined milestones for approved draft and approved publication.",
        },
        {
            "title": "4.5 Beneficiary interview or field story",
            "deliverable": "Record and publish an interview or field story based on a campaign brief. The task must state the consent process, approved questions, privacy restrictions and whether the subject's identity may be shown.",
            "proof": "Consent confirmation, final content file, public URL when publication is required, supporting photographs and analytics where applicable.",
            "evaluation": "Consent documented; facts and quotations handled responsibly; privacy conditions followed; required campaign message included; content approved and published as instructed.",
            "payment": "Payment after consent, content and publication requirements have all been accepted.",
        },
    ]

    for example in examples:
        doc.add_heading(example["title"], level=2)
        add_label_paragraph(doc, "Deliverable", example["deliverable"], keep_with_next=True)
        add_label_paragraph(doc, "Possible proof", example["proof"], keep_with_next=True)
        add_label_paragraph(doc, "Evaluation", example["evaluation"], keep_with_next=True)
        add_label_paragraph(doc, "Payment basis", example["payment"])

    add_note(
        doc,
        "Current scope boundary",
        "This version of the concept note ends with representative task formats. Detailed evaluation scorecards, dashboard report layouts, payment integration options, governance policies and implementation phases can be added after the operating model is approved.",
    )


def add_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Government Influencer Campaign Management Platform"
    props.subject = "Product concept, operational flow, hosting resources and example task formats"
    props.author = "Project Team"
    props.keywords = "government, influencer, campaign, proof of work, evaluation, payment tracking"
    props.comments = "Working concept note"
    props.created = datetime(2026, 8, 6, 0, 0, 0)


def build() -> None:
    doc = Document()
    style_document(doc)
    configure_header_footer(doc.sections[0])
    bullet_num, decimal_num = add_numbering_definitions(doc)
    add_core_properties(doc)

    add_title_page(doc)
    add_platform_overview(doc, bullet_num)
    add_users(doc)
    add_flow(doc, decimal_num)
    add_task_examples(doc, bullet_num)

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
